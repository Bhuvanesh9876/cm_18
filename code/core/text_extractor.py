"""
Text Extractor Module

This module handles the extraction of text from PDF and DOCX files.
It employs multiple strategies (PyMuPDF, pdfplumber, pdfminer) to ensure
robust text extraction from various PDF formats.
"""

import io
import re
from typing import Any
import fitz
import pdfplumber
import chardet
from pdfminer.high_level import extract_text as pdfminer_extract
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_bytes
except ImportError:
    Image = None
    pytesseract = None
    convert_from_bytes = None

def _extract_with_pymupdf(file: Any) -> str:
    """Extract text using PyMuPDF (fitz)."""
    try:
        if hasattr(file, "read"):
            data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                data = f.read()

        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                pages.append(text.strip())
        doc.close()
        return "\n".join(pages)
    except Exception:
        return ""

def _extract_with_pdfplumber(file: Any) -> str:
    """Extract text using pdfplumber."""
    try:
        if hasattr(file, "seek"):
            file.seek(0)
        with pdfplumber.open(file) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
        return "\n".join(pages)
    except Exception:
        return ""

def _extract_with_pdfminer(file: Any) -> str:
    """Extract text using pdfminer.six."""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
            return pdfminer_extract(io.BytesIO(data))
        return pdfminer_extract(file)
    except Exception:
        return ""

def _extract_from_docx(file: Any) -> str:
    """Extract text from DOCX files using python-docx with zipfile XML fallback.

    Strategy:
    1. Standard python-docx (paragraphs + tables).
    2. Raw XML extraction via zipfile — handles corrupted/non-standard DOCX.
    3. Returns detailed error on total failure (for server-side logging).
    """
    import logging
    logger = logging.getLogger(__name__)

    # Read bytes once so we can retry with different strategies
    try:
        if hasattr(file, "read"):
            file.seek(0)
            raw_bytes = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                raw_bytes = f.read()
    except Exception as e:
        logger.warning(f"DOCX: failed to read file bytes: {e}")
        return ""

    # ---- Strategy 1: python-docx (paragraphs + tables) ----
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw_bytes))
        parts: list[str] = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = []
                seen: set[str] = set()
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct and ct not in seen:
                        seen.add(ct)
                        cells.append(ct)
                if cells:
                    parts.append("  ".join(cells))

        if parts:
            return "\n".join(parts)
        logger.info("DOCX strategy 1 returned empty; trying zipfile fallback.")
    except Exception as e:
        logger.warning(f"DOCX strategy 1 (python-docx) failed: {e}")

    # ---- Strategy 2: zipfile raw XML extraction (for corrupted DOCX) ----
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        parts = []
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
            # word/document.xml holds the main body
            xml_names = [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")]
            for xml_name in xml_names:
                with z.open(xml_name) as xf:
                    tree = ET.parse(xf)
                    root = tree.getroot()
                    for elem in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                        text = elem.text
                        if text and text.strip():
                            parts.append(text.strip())

        if parts:
            logger.info("DOCX strategy 2 (zipfile XML) succeeded.")
            return "\n".join(parts)
        logger.warning("DOCX strategy 2 (zipfile XML) returned empty text.")
    except Exception as e:
        logger.warning(f"DOCX strategy 2 (zipfile XML) failed: {e}")

    return ""


def _extract_from_txt(file: Any) -> str:
    """Extract text from TXT files with encoding detection."""
    try:
        if hasattr(file, "read"):
            raw_data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                raw_data = f.read()
        
        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'
        return raw_data.decode(encoding)
    except Exception:
        return ""

def _extract_with_ocr(file: Any) -> str:
    """Extract text from images using OCR (pytesseract)."""
    if not pytesseract or not Image:
        return ""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
            img = Image.open(io.BytesIO(data))
        else:
            img = Image.open(file)
        
        text = pytesseract.image_to_string(img)
        return text
    except Exception:
        return ""

def _extract_scanned_pdf(file: Any) -> str:
    """Fallback for scanned PDFs: convert to images and then OCR."""
    if not convert_from_bytes or not pytesseract:
        return ""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                data = f.read()
        
        images = convert_from_bytes(data)
        full_text = []
        for img in images:
            text = pytesseract.image_to_string(img)
            full_text.append(text)
        return "\n".join(full_text)
    except Exception:
        return ""

def _post_clean(text: str) -> str:
    """
    Clean and normalize extracted text.
    Removes invisible characters, weird hyphenation, and excessive newlines.
    """
    if not text:
        return ""
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.replace("\x00", "")
    return text.strip()

def extract_text(file: Any) -> str:
    """
    Main entry point for text extraction.
    Handles PDF, DOCX, TXT, and Images using multiple strategies.
    """
    filename = ""
    if hasattr(file, "name"):
        filename = file.name.lower()

    # Handle DOCX
    if filename.endswith(".docx"):
        text = _extract_from_docx(file)
        return _post_clean(text)
    
    # Handle TXT
    if filename.endswith(".txt"):
        text = _extract_from_txt(file)
        return _post_clean(text)

    # Handle Images
    if filename.endswith((".png", ".jpg", ".jpeg")):
        text = _extract_with_ocr(file)
        return _post_clean(text)

    # Handle PDF with tiered extraction + OCR Fallback
    # 1. PyMuPDF (Fast)
    text = _extract_with_pymupdf(file)
    if len(text.strip()) > 100:
        return _post_clean(text)

    # 2. pdfplumber (Better layout)
    text = _extract_with_pdfplumber(file)
    if len(text.strip()) > 100:
        return _post_clean(text)

    # 3. pdfminer (Fallback)
    text = _extract_with_pdfminer(file)
    if len(text.strip()) > 100:
        return _post_clean(text)
    
    # 4. OCR Fallback for Scanned PDFs
    text = _extract_scanned_pdf(file)
    return _post_clean(text)
